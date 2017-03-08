from docx import Document
from tkFileDialog import askopenfilename

#For the script to work perfectly, pdf transcripts have to be converted to docx using https://online2pdf.com/pdf2docx
#This converter formats the document in such a way that complies with the 'collectPoints' function 
docFile = askopenfilename()

#Identifies the column containing the marks for each completed subject throughout the degree 
def collectPoints():
    document = Document(docFile)
    table = 0
    marksList = 0
    sems = []
    marks = {}
    marks[marksList] = []

    while table < len(document.tables):
        markCol = document.tables[table].columns[5].cells
        markCel = 0
        while markCel < len(markCol):
            try:
                val = int(markCol[markCel].paragraphs[0].text)
                marks[marksList].extend([val])
                if len(marks[marksList]) == 4:
                    sems.append(marks[marksList])
                    marksList += 1
                    marks[marksList] = []
                    #Separate marks into their respective semesters
            except ValueError:
                pass
                #Ignore anything that cannot be converted into an integer
            markCel +=1
        table +=1 
        #If the document contains multiple pages, additional tables 
        #containing additional marks may be formed. 

    return sems

